import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime, timezone, timedelta
from docxtpl import DocxTemplate
import re
import json
import os

st.set_page_config(page_title="COOR TOOL VJ DAD", layout="wide")

# Lấy giờ Việt Nam
now_vn = datetime.now(timezone(timedelta(hours=7)))

# ── helpers ──────────────────────────────────────────────────────────────────

def _vn_now():
    return datetime.now(timezone(timedelta(hours=7)))

# ── STREAMLIT TAB BÁO CÁO SỰ VỤ ───────────────────────────────────────────────

def tab_bao_cao_su_vu():
    st.subheader("📋 Trình tạo Báo cáo Sự cố Kỹ thuật Tàu bay (CAAV)")
    st.caption("Nhập thông tin vào các ô bên dưới để điền vào file template (template_bao_cao.docx).")

    now = _vn_now()

    with st.form("form_bao_cao"):

        st.markdown("##### 📅 Ngày tạo báo cáo (ngay_thang_nam)")
        c1, c2, c3 = st.columns(3)
        ngay  = c1.text_input("Ngày",  value=str(now.day).zfill(2))
        thang = c2.text_input("Tháng", value=str(now.month).zfill(2))
        nam   = c3.text_input("Năm",   value=str(now.year))
        
        # Biến ngày tháng năm đầy đủ để điền vào template
        ngay_thang_nam = f"{ngay} tháng {thang} năm {nam}"
        st.divider()
        st.markdown("##### ✈️ 1. Thông tin tổng quát")

        c1, c2 = st.columns(2)
        loai_may_bay = c1.text_input("Loại máy bay (loai_tau):", value="A321")
        so_hieu_tau  = c2.text_input("Số hiệu tàu bay (reg):", placeholder="VD: VN-A677")

        c1, c2 = st.columns(2)
        chuyen_bay  = c1.text_input("Số hiệu chuyến bay (flight_no):", placeholder="VD: VJ517")
        noi_di_den  = c2.text_input("Nơi đi – Nơi đến (route):", placeholder="VD: HAN-DAD")

        c1, c2 = st.columns(2)
        gio_su_co  = c1.text_input("Giờ xảy ra sự cố (LT):", placeholder="VD: 20:45")
        ngay_su_co = c2.text_input("Ngày xảy ra sự cố:", placeholder="VD: 31/03/2026",
                                   value=now.strftime("%d/%m/%Y"))
        
        # Timing tổng hợp (nếu template dùng chung)
        timing_combined = f"{gio_su_co} LT, ngày {ngay_su_co}"

        st.divider()
        st.markdown("##### 🔍 2. Mô tả sự cố kỹ thuật (description)")
        mo_ta_su_co = st.text_area("Nội dung sự cố", placeholder="VD: TAILSTRIKE DURING GO-AROUND", height=80)

        st.divider()
        st.markdown("##### 🔧 3. Khắc phục sự cố (action)")
        khac_phuc = st.text_area(
            "Các bước đã thực hiện (mỗi dòng một bước)",
            placeholder=(
                "- VJ517 ĐÁP VỀ BÃI 25.\n"
                "- A677 NẰM SÂN, ĐỔI TÀU BAY KHÁC BAY TIẾP CHUYẾN VJ520.\n"
                "- THỰC HIỆN KÉO A677 VỀ BÃI 5M ĐỂ TIẾP TỤC ĐIỀU TRA VÀ KHẮC PHỤC SỰ CỐ."
            ),
            height=150,
        )

        st.divider()
        st.markdown("##### 🖊️ Người báo cáo (reporter)")
        nguoi_bao_cao = st.text_input("Họ tên người báo cáo (Ký tên)", placeholder="VD: PHAM VIET DUC")

        submitted = st.form_submit_button("📄 Điền template & Tải xuống (.docx)",
                                          use_container_width=True, type="primary")

    if submitted:
        # Kiểm tra file template
        template_file = "template_bao_cao.docx"
        if not os.path.exists(template_file):
            st.error(f"❌ Không tìm thấy file mẫu: `{template_file}`. Vui lòng tải file này lên thư mục dự án.")
        else:
            try:
                # Load template
                doc = DocxTemplate(template_file)
                
                # Chuẩn bị dữ liệu khớp với biến trong file Word của bạn
                context = {
                    "ngay_thang_nam": ngay_thang_nam,
                    "ngay": ngay,
                    "thang": thang,
                    "nam": nam,
                    "loai_tau": loai_may_bay.upper(),
                    "reg": so_hieu_tau.upper() if "VN-" in so_hieu_tau.upper() else f"VN-{so_hieu_tau.upper()}",
                    "flight_no": chuyen_bay.upper(),
                    "route": noi_di_den.upper(),
                    "timing": timing_combined,
                    "description": mo_ta_su_co.upper(),
                    "action": khac_phuc.upper(),
                    "reporter": nguoi_bao_cao.upper()
                }
                
                # Render (điền dữ liệu)
                doc.render(context)
                
                # Xuất ra buffer
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                ten_file_xuat = f"BAO_CAO_SU_CO_{so_hieu_tau.upper()}_{datetime.now().strftime('%d%m')}.docx"

                st.success("✅ Đã điền báo cáo vào template thành công!")
                st.download_button(
                    label="📥 Tải Báo Cáo (.docx)",
                    data=buffer,
                    file_name=ten_file_xuat,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"⚠️ Lỗi khi xử lý template: {e}")

            # Preview nhanh bằng HTML
            with st.expander("👁️ Xem trước nội dung (Mô phỏng)"):
                st.markdown(f"""
<div style="font-family: 'Times New Roman', Times, serif; font-size: 13pt; border: 1px solid #ddd; padding: 40px; background-color: white; color: black; max-width: 800px; margin: auto; line-height: 1.15;">
    <div style="text-align: right; font-size: 11pt; font-family: Calibri, sans-serif;">
        Đà Nẵng, ngày {ngay_thang_nam}
    </div>
    <br><br>
    <div style="text-align: center; font-weight: bold; font-size: 14pt;">BÁO CÁO SỰ CỐ KỸ THUẬT TÀU BAY</div>
    <br><br>
    <div><b>1. THÔNG TIN TỔNG QUÁT:</b></div>
    <div style="margin-bottom: 2px;">- Loại máy bay: {loai_may_bay.upper()}</div>
    <div style="margin-bottom: 2px;">- Số hiệu tàu bay: VN-{so_hieu_tau.upper()}</div>
    <div style="margin-bottom: 2px;">- Số hiệu chuyến bay: {chuyen_bay.upper()}</div>
    <div style="margin-bottom: 2px;">- Nơi đi – Nơi đến: {noi_di_den.upper()}</div>
    <div style="margin-bottom: 2px;">- Giờ, ngày xảy ra sự cố: {timing_combined}</div>
    <br>
    <div><b>2. MÔ TẢ SỰ CỐ KỸ THUẬT:</b></div>
    <div>{mo_ta_su_co.upper()}</div>
    <br>
    <div><b>3. KHẮC PHỤC SỰ CỐ:</b></div>
    {chr(10).join([f'<div style="margin-bottom: 2px;">- {d.strip().lstrip("- ").upper()}</div>' for d in khac_phuc.split(chr(10)) if d.strip()])}
    <br><br><br>
    <div style="text-align: right; margin-right: 50px;">NGƯỜI BÁO CÁO</div>
    <br><br><br>
    <div style="text-align: right; margin-right: 50px;"><b>{nguoi_bao_cao.upper()}</b></div>
</div>
""", unsafe_allow_html=True)

tab1, tab2, tab3, tab4 = st.tabs(["✈️ Kế hoạch Kéo tàu", "📋 Báo cáo Sự cố CAAV", "🔋 Request SAGS GPU", "📝 Request ONE_OFF"])

with tab1:
    st.title("✈️ Trình tạo mail kéo tàu Vietjet DAD")
    st.caption(f"Ngày tạo báo cáo: {now_vn.strftime('%d/%m/%Y')}")
    # ---HÀM JSON---
    DATA_FILE = "plans_data.json"

    def save_plans(plans):
        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(plans, f, ensure_ascii=False, indent=2)

    def load_plans():
        if os.path.exists(DATA_FILE):
            with open(DATA_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        return []
    # --- KHỞI TẠO SESSION STATE ---
    if 'plans' not in st.session_state:
        st.session_state.plans = load_plans()
    if 'editing_index' not in st.session_state:
        st.session_state.editing_index = None

    # --- HÀM HỖ TRỢ ---
    def create_word_document(content):
        doc = Document()
        # Loại bỏ các tag HTML highlight và ký hiệu == nếu có trước khi xuất Word
        clean_content = re.sub(r'<span.*?>|<\/span>', '', content)
        clean_content = clean_content.replace('==', '')
        bold_pattern = re.compile(r'\*\*(.*?)\*\*')
        for line in clean_content.split('\n'):
            p = doc.add_paragraph()
            p.style.font.name = 'Arial'
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if bold_pattern.match(part):
                    p.add_run(part.strip('*')).bold = True
                else:
                    p.add_run(part)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    def generate_report_content(plans, highlight=False, kinh_gui=""):
        today_str = now_vn.strftime('%d/%m/%Y')
        
        # Header 
        report_lines = []
        # Thêm tiêu đề mail
        report_lines.append(f"**KẾ HOẠCH KÉO TÀU VJ DAD NGÀY {today_str}**")
        report_lines.append("")  # Dòng trống sau tiêu đề
        if kinh_gui:
            report_lines.extend(kinh_gui.split('\n'))
        
        report_lines.append("") # Dòng trống sau Kính gửi
        report_lines.append(f"VJ DAD gửi kế hoạch kéo/đẩy tàu bay ngày {today_str} như sau:")
        report_lines.append("")

        # Body
        for i, plan in enumerate(plans):
            changed = plan.get('changed_fields', [])
            
            # Tiêu đề mục (Ưu tiên hiển thị Đang bãi)
            tau_str = f"VN-{plan['Tàu']}"
            if 'Tàu' in changed and highlight: tau_str = f"=={tau_str}=="
            
            if plan.get('Đang bãi'):
                db_str = f"Đang bãi {plan['Đang bãi']}"
                if 'Đang bãi' in changed and highlight: db_str = f"=={db_str}=="
                title = f"**{i+1}. {tau_str}/ {db_str}**"
            else:
                ch_str = plan['Chuyến']
                if 'Chuyến' in changed and highlight: ch_str = f"=={ch_str}=="
                title = f"**{i+1}. {tau_str}/{ch_str}**"
                if plan['STA']:
                    sta_str = f"STA {plan['STA']}"
                    if 'STA' in changed and highlight: sta_str = f"=={sta_str}=="
                    title += f" **{sta_str}**"
            
            if plan['Ghi chú']:
                gc_upper = str(plan['Ghi chú']).upper()
                if "CNX" in gc_upper:
                    gc_display = "HUỶ KẾ HOẠCH KÉO"
                elif "DONE" in gc_upper:
                    gc_display = "ĐÃ HOÀN THÀNH"
                else:
                    gc_display = plan['Ghi chú']
                gc_str = f"({gc_display})"
                if 'Ghi chú' in changed and highlight: gc_str = f"=={gc_str}=="
                title += f" **{gc_str}**"
            report_lines.append(title)
            # Kiểm tra điều kiện ẩn chi tiết
            gc_upper = str(plan['Ghi chú']).upper()
            if "CNX" in gc_upper or "DONE" in gc_upper:
                report_lines.append("")
                continue

            # Kéo tới (Luôn hiện nếu có thông tin)
            if plan['Kéo tới']:
                kt_val = plan['Kéo tới']
                if 'Kéo tới' in changed and highlight: kt_val = f"=={kt_val}=="
                line = f"    - Kéo về **{kt_val}**"
                
                if plan['Thời gian kéo']:
                    tgk_val = plan['Thời gian kéo']
                    if 'Thời gian kéo' in changed and highlight: tgk_val = f"=={tgk_val}=="
                    line += f" dự kiến vào lúc: **{tgk_val}**"
                report_lines.append(line)

            if plan['Kéo ga lớn'] == "CÓ":
                line = f"    - Kéo ra ga lớn: **CÓ**"
                if plan['Thời gian kéo ga lớn']:
                    tggl_val = plan['Thời gian kéo ga lớn']
                    if 'Thời gian kéo ga lớn' in changed and highlight: tggl_val = f"=={tggl_val}=="
                    line += f". Thời gian dự kiến: **{tggl_val}**"
                report_lines.append(line)
            
            if plan['Kéo khai thác'] == "CÓ":
                if not plan['Thời gian kéo khai thác'] or plan['Thời gian kéo khai thác'] == "THÔNG BÁO SAU":
                    line = "    - Kéo ra bãi khai thác: **CÓ**. Thời gian dự kiến: **THÔNG BÁO SAU**"
                else:
                    ktch_val = plan['Khai thác chuyến']
                    if 'Khai thác chuyến' in changed and highlight: ktch_val = f"=={ktch_val}=="
                    tgkt_val = plan['Thời gian kéo khai thác']
                    if 'Thời gian kéo khai thác' in changed and highlight: tgkt_val = f"=={tgkt_val}=="
                    line = f"    - Kéo ra bãi khai thác chuyến: **{ktch_val}**. Thời gian dự kiến: **{tgkt_val}**"
                report_lines.append(line)

            dv_val = plan['Đơn vị kéo']
            if 'Đơn vị kéo' in changed and highlight: dv_val = f"=={dv_val}=="
            report_lines.append(f"    - Đơn vị kéo đẩy: **{dv_val}**")
            
            asu_val = plan['ASU-GPU']
            if 'ASU-GPU' in changed and highlight: asu_val = f"=={asu_val}=="
            report_lines.append(f"    - Cần ASU, GPU: **{asu_val}**")
            report_lines.append("") # Dòng trống giữa các mục

        # Footer cố định
        report_lines.append("Kính mong TBT, ĐHSĐ sắp xếp tàu về bến thuận tiện cho việc kéo đẩy.")
        report_lines.append("VJ sẽ cập nhật thông tin khi kế hoạch thay đổi.")
        
        return "\n".join(report_lines)

    def convert_to_html(markdown_text):
        """Chuyển đổi Markdown sang HTML để hỗ trợ copy paste bôi đen giữ định dạng."""
        html = markdown_text.replace('\n', '<br>')
        # Xử lý highlight ==text== -> <span style="background-color: yellow">text</span>
        html = re.sub(r'==(.*?)==', r'<span style="background-color: #FFFF00; color: black; padding: 0 2px; border-radius: 2px;">\1</span>', html)
        # Xử lý in đậm **text** -> <b>text</b>
        html = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', html)
        # Xử lý thụt lề (4 dấu cách) -> &nbsp;
        html = html.replace('    ', '&nbsp;&nbsp;&nbsp;&nbsp;')
        return f'<div style="font-family: Arial; font-size: 14px; line-height: 1.5; color: black;">{html}</div>'

    # --- SIDEBAR - Cấu hình ---
    st.sidebar.header("⚙️ Cấu hình Mail mẫu")
    default_kinh_gui = """Kính Gửi:
        -Trực Ban Trưởng
        -Điều Hành Sân Đỗ
        -Đài kiểm soát mặt đất"""

    kinh_gui_input = st.sidebar.text_area(
        "Danh sách Kính gửi (Edit tại đây):", 
        value=default_kinh_gui,
        height=150
    )

    st.sidebar.divider()
    st.sidebar.info("Footer đã được cố định theo mẫu.")

    # --- FORM NHẬP LIỆU ---
    st.subheader("📋 1. Nhập Kế hoạch Kéo tàu")

    # Lấy dữ liệu nếu đang ở chế độ chỉnh sửa
    edit_idx = st.session_state.editing_index
    edit_data = st.session_state.plans[edit_idx] if edit_idx is not None else {}

    with st.form("plan_form", clear_on_submit=True):
        if edit_idx is not None:
            st.warning(f"🔄 Đang chỉnh sửa kế hoạch số {edit_idx + 1}")
        
        st.write("Điền thông tin cho một tàu bay và nhấn nút Thêm/Cập nhật.")
        c1, c2, c3, c4, c5 = st.columns(5)
        tau = c1.text_input("Tàu (VN-)", value=edit_data.get("Tàu", ""), placeholder="A662")
        chuyen = c2.text_input("Chuyến", value=edit_data.get("Chuyến", ""), placeholder="VJ703")
        sta = c3.text_input("STA / Ghi chú", value=edit_data.get("STA", ""), placeholder="12:30 hoặc PHASE CHECK")
        dang_bai = c4.text_input("Đang bãi", value=edit_data.get("Đang bãi", ""), placeholder="VJ01 hoặc 3M")
        ghi_chu = c5.text_input("Ghi chú thêm (nếu có)", value=edit_data.get("Ghi chú", ""))

        st.write("Kéo về bãi:")
        c1, c2 = st.columns([1, 2])
        keo_toi = c1.text_input("Kéo về bãi", value=edit_data.get("Kéo tới", ""), placeholder="VJ01")
        tg_keo = c2.text_input("Thời gian kéo về bãi", value=edit_data.get("Thời gian kéo", ""), placeholder="11:00")

        st.write("Kéo ra ga lớn & khai thác:")
        c1, c2, c3, c4, c5 = st.columns(5)
        
        kg_idx = ["KHÔNG", "CÓ"].index(edit_data.get("Kéo ga lớn", "KHÔNG"))
        keo_ga_lon = c1.selectbox("Kéo ga lớn?", ["KHÔNG", "CÓ"], index=kg_idx)
        tg_ga_lon = c2.text_input("Giờ kéo ga lớn", value=edit_data.get("Thời gian kéo ga lớn", "THÔNG BÁO SAU"))
        
        kk_idx = ["CÓ", "KHÔNG"].index(edit_data.get("Kéo khai thác", "CÓ"))
        keo_kt = c3.selectbox("Kéo khai thác?", ["CÓ", "KHÔNG"], index=kk_idx)
        kt_chuyen = c4.text_input("Chuyến khai thác", value=edit_data.get("Khai thác chuyến", ""), placeholder="VJ703")
        tg_kt = c5.text_input("Giờ kéo khai thác", value=edit_data.get("Thời gian kéo khai thác", "THÔNG BÁO SAU"))

        st.write("Thông tin khác:")
        c1, c2 = st.columns(2)
        don_vi = c1.text_input("Đơn vị kéo", value=edit_data.get("Đơn vị kéo", "VJ"))
        
        asu_idx = ["KHÔNG", "CÓ"].index(edit_data.get("ASU-GPU", "KHÔNG"))
        asu_gpu = c2.selectbox("Cần ASU/GPU?", ["KHÔNG", "CÓ"], index=asu_idx)

        btn_label = "➕ Thêm vào danh sách" if edit_idx is None else "💾 Cập nhật kế hoạch"
        submitted = st.form_submit_button(btn_label, use_container_width=True)
        
        if submitted and tau:
            new_plan = {
                "Tàu": tau, "Chuyến": chuyen, "STA": sta, "Đang bãi": dang_bai, "Ghi chú": ghi_chu,
                "Kéo tới": keo_toi, "Thời gian kéo": tg_keo,
                "Kéo ga lớn": keo_ga_lon, "Thời gian kéo ga lớn": tg_ga_lon,
                "Kéo khai thác": keo_kt, "Khai thác chuyến": kt_chuyen, "Thời gian kéo khai thác": tg_kt,
                "Đơn vị kéo": don_vi, "ASU-GPU": asu_gpu
            }
            
            if edit_idx is not None:
                # So sánh để tìm các trường thay đổi
                old_plan = st.session_state.plans[edit_idx]
                changed_fields = []
                for k, v in new_plan.items():
                    if v != old_plan.get(k):
                        changed_fields.append(k)
                
                new_plan['changed_fields'] = changed_fields
                st.session_state.plans[edit_idx] = new_plan
                save_plans(st.session_state.plans)
                st.session_state.editing_index = None
                st.success("Đã cập nhật kế hoạch!")
            else:
                new_plan['changed_fields'] = []
                st.session_state.plans.append(new_plan)
                save_plans(st.session_state.plans)
                st.success("Đã thêm kế hoạch mới!")
            st.rerun()

    if st.session_state.editing_index is not None:
        if st.button("❌ Hủy chỉnh sửa", use_container_width=True):
            st.session_state.editing_index = None
            st.rerun()

    # --- DANH SÁCH KẾ HOẠCH ĐÃ THÊM ---
    st.subheader("📝 2. Danh sách Kế hoạch")
    if not st.session_state.plans:
        st.info("Chưa có kế hoạch nào được thêm.")
    else:
        for i, plan in enumerate(st.session_state.plans):
            with st.container(border=True):
                c1, c2, c3, c4, c5 = st.columns([6, 1, 1, 1, 1])
                
                # Hiển thị thông tin tiêu đề trong danh sách
                if plan.get('Đang bãi'):
                    info = f"**{i+1}. VN-{plan['Tàu']}/ Đang bãi {plan['Đang bãi']}**"
                else:
                    info = f"**{i+1}. VN-{plan['Tàu']}/{plan['Chuyến']}** - STA: {plan['STA']}"
                
                if plan['Ghi chú']: info += f" ({plan['Ghi chú']})"
                
                details = []
                if plan['Kéo tới']: 
                    details.append(f"Kéo về: {plan['Kéo tới']} ({plan['Thời gian kéo']})")
                if plan['Kéo ga lớn'] == "CÓ": details.append(f"Ga lớn: CÓ ({plan['Thời gian kéo ga lớn']})")
                if plan['Kéo khai thác'] == "CÓ": 
                    if not plan['Thời gian kéo khai thác'] or plan['Thời gian kéo khai thác'] == "THÔNG BÁO SAU":
                        details.append("Khai thác: CÓ (THÔNG BÁO SAU)")
                    else: details.append(f"Khai thác: {plan['Khai thác chuyến']} ({plan['Thời gian kéo khai thác']})")
                
                c1.markdown(info)
                if details:
                    c1.caption(" | ".join(details))
                
                # Nút di chuyển lên
                if c2.button("🔼", key=f"up_{i}", help="Di chuyển lên"):
                    if i > 0:
                        st.session_state.plans[i], st.session_state.plans[i-1] = st.session_state.plans[i-1], st.session_state.plans[i]
                        save_plans(st.session_state.plans)
                        st.rerun()
                
                # Nút di chuyển xuống
                if c3.button("🔽", key=f"down_{i}", help="Di chuyển xuống"):
                    if i < len(st.session_state.plans) - 1:
                        st.session_state.plans[i], st.session_state.plans[i+1] = st.session_state.plans[i+1], st.session_state.plans[i]
                        st.rerun()

                # Nút sửa
                if c4.button("📝", key=f"edit_{i}", help="Chỉnh sửa"):
                    st.session_state.editing_index = i
                    st.rerun()
                
                # Nút xóa
                if c5.button("❌", key=f"del_{i}", help="Xóa"):
                    st.session_state.plans.pop(i)
                    save_plans(st.session_state.plans)
                    if st.session_state.editing_index == i:
                        st.session_state.editing_index = None
                    st.rerun()

    # --- TẠO MAIL ---
    st.divider()
    col_btn1, col_btn2 = st.columns(2)
    show_highlight = col_btn1.toggle("Highlight thông tin thay đổi", value=True)

    if st.button("🚀 3. Tạo Mail mẫu", use_container_width=True, type="primary"):
        if not st.session_state.plans:
            st.warning("Vui lòng thêm ít nhất một kế hoạch trước khi tạo mail.")
        else:
            st.subheader("📧 4. Kết quả Mail mẫu")
            report_content = generate_report_content(
                st.session_state.plans, 
                highlight=show_highlight, 
                kinh_gui=kinh_gui_input
            )
            
            # Hiển thị HTML để người dùng có thể bôi đen và copy paste giữ định dạng
            st.caption("👇 Bạn có thể bôi đen (Highlight) nội dung dưới đây để Copy & Paste vào Mail/Zalo (giữ nguyên định dạng in đậm):")
            st.write(convert_to_html(report_content), unsafe_allow_html=True)
            
            st.divider()
            st.caption("Hoặc copy text thô tại đây (Không bao gồm Highlight):")
            # Text thô không nên có highlight markdown ==text==
            raw_content = re.sub(r'==(.*?)==', r'\1', report_content)
            st.code(raw_content, language="text")
            
            st.download_button(
                label="📄 Tải xuống file Word (.docx)",
                data=create_word_document(report_content),
                file_name=f"Ke_hoach_keo_tau_{now_vn.strftime('%d%m%Y')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
with tab2:
    tab_bao_cao_su_vu()

with tab3:
    st.title("🔋 Request SAGS phục vụ GPU")
    st.caption(f"Ngày hiện tại: {now_vn.strftime('%d/%m/%Y')}")

    # --- HÀM TRÍCH XUẤT DỮ LIỆU ---
    def parse_sags_gpu_input(text):
        import re
        lines = text.strip().split('\n')
        results = []
        for line in lines:
            if not line.strip(): continue
            
            # Tách các thành phần bằng khoảng trắng hoặc tab
            parts = line.split()
            if len(parts) < 8: continue
            
            # Tìm số hiệu tàu (thường bắt đầu bằng VN- hoặc HS-)
            ac_reg = ""
            for p in parts:
                if p.startswith("VN-") or p.startswith("HS-"):
                    ac_reg = p
                    break
            
            # Tìm các mốc thời gian HH:MM
            times = re.findall(r'\d{2}:\d{2}', line)
            
            try:
                # Heuristic trích xuất dựa trên cấu trúc copy từ web:
                # # Date Flight Type Route AC_Reg ... ARR ... DEP
                # Example: 14 02-Apr VJ721 / VJ633 S HPH-DAD / DAD-SGN VN-A648 ... 12:35 ... 14:50
                
                # Flight: Lấy phần trước dấu gạch chéo đầu tiên
                flt_no = parts[2]
                
                # Route: Lấy phần trước dấu gạch chéo thứ hai
                # Tìm index của dấu gạch chéo đầu tiên trong parts sau index 2
                slash_indices = [i for i, x in enumerate(parts) if x == "/"]
                
                route = ""
                if len(slash_indices) >= 2:
                    route = parts[slash_indices[1]-1] # Phần trước dấu / thứ 2
                else:
                    # Fallback nếu không tìm thấy cấu trúc chuẩn
                    for p in parts:
                        if "-" in p and not p.startswith("VN-") and not p.startswith("HS-"):
                            route = p
                            break
                
                sta = times[0] if len(times) >= 1 else ""
                std = times[-1] if len(times) >= 2 else ""
                
                results.append({
                    "FLIGHT": f"{flt_no} / ____",
                    "ROUTE": f"{route} / ____",
                    "A/C": ac_reg,
                    "STA": sta,
                    "STD": "" # Theo mẫu ảnh 1 STD để trống
                })
            except:
                continue
        return results

    # --- UI TAB 3 ---
    st.subheader("1. Dán dữ liệu từ lịch bay")
    st.info("Copy nguyên dòng từ web lịch bay (có đủ cột #, Date, Flight, Route...) rồi dán vào ô dưới đây. Mỗi chuyến một dòng.")
    
    raw_data = st.text_area("Dán dữ liệu vào đây:", height=200, placeholder="14 02-Apr VJ721 / VJ633 S HPH-DAD / DAD-SGN VN-A648 A321 12:35 12:35 14:50")
    
    if raw_data:
        parsed_list = parse_sags_gpu_input(raw_data)
        
        if not parsed_list:
            st.error("Không thể trích xuất dữ liệu. Vui lòng kiểm tra lại định dạng dán vào.")
        else:
            st.subheader("2. Kiểm tra dữ liệu đã trích xuất")
            st.table(parsed_list)
            
            # --- TẠO NỘI DUNG MAIL ---
            date_str = now_vn.strftime('%d/%m/%Y')
            subject = f"DANH SÁCH TÀU DỰ KIẾN CẦN SAGS PHỤC VỤ GPU - NGÀY {date_str}"
            
            mail_body_top = f"""Dear Sags,
Theo yêu cầu của Sags, để tiện việc bố trí nhân sự vận hành GPU của Sags.
VJ DAD gửi danh sách tàu có kế hoạch bảo dưỡng cần SAGS phục vụ GPU (trong trường hợp GPU của VJ không đủ) ngày {date_str} như sau:
"""
            mail_body_bottom = "VJ DAD sẽ update lại nếu có sự thay đổi!"

            # Tạo bảng HTML cho mail
            table_html = """<table border="1" style="border-collapse: collapse; width: auto; min-width: 600px; font-family: Arial; font-size: 13px;">
<tr style="background-color: #f2f2f2; text-align: center;">
<th style="padding: 5px 10px;">STT</th>
<th style="padding: 5px 10px;">DATE</th>
<th style="padding: 5px 15px;">FLIGHT</th>
<th style="padding: 5px 20px;">ROUTE</th>
<th style="padding: 5px 15px;">A/C</th>
<th style="padding: 5px 10px;">STA</th>
<th style="padding: 5px 10px;">STD</th>
<th style="padding: 5px 20px;">NOTE</th>
</tr>"""
            date_short = now_vn.strftime('%d/%m/%y')
            for i, item in enumerate(parsed_list):
                table_html += f"""<tr style="text-align: center;">
<td style="padding: 5px;">{i+1}</td>
<td style="padding: 5px;">{date_short}</td>
<td style="padding: 5px;">{item['FLIGHT']}</td>
<td style="padding: 5px;">{item['ROUTE']}</td>
<td style="padding: 5px;">{item['A/C']}</td>
<td style="padding: 5px;">{item['STA']}</td>
<td style="padding: 5px;">{item['STD']}</td>
<td style="padding: 5px;"></td>
</tr>"""
            table_html += "</table>"

            st.divider()
            st.subheader("3. Mẫu Email (Sẵn sàng để Copy)")
            
            # Hiển thị Tiêu đề
            st.markdown(f"**Tiêu đề Mail:**")
            st.code(subject, language="text")
            
            # Hiển thị Nội dung (HTML Preview)
            st.markdown(f"**Nội dung Mail:**")
            full_html = f"""<div style="font-family: Arial; font-size: 14px; color: black;">
{mail_body_top.replace('\n', '<br>')}
<br>
{table_html}
<br>
{mail_body_bottom}
</div>"""
            st.markdown(full_html, unsafe_allow_html=True)
            
            st.caption("👇 Bôi đen toàn bộ nội dung trên (bao gồm cả bảng) để Copy & Paste vào Outlook/Gmail.")

with tab4:
    st.title("📝 Request ONE-OFF AUTHORIZATION")
    st.caption(f"Ngày tạo: {now_vn.strftime('%d/%m/%Y')}")

    with st.form("form_one_off"):
        st.markdown("##### 👤 1. Applicant Information")
        c1, c2 = st.columns(2)
        applicant_name = c1.text_input("Applicant name:", placeholder="VD: NGUYEN HUY HOANG")
        auth_no = c2.text_input("Auth. No.:", placeholder="VD: VJC.CRS.439")
        
        c1, c2 = st.columns(2)
        licence_no = c1.text_input("Licence Number:", placeholder="VD: 54361-AMT")
        expiry_date = c2.text_input("Expiry date:", placeholder="VD: 31MAY2029")
        
        auth_scope = st.text_input("Current authorization scope:", value="A320/A321 CAT A ( CFM56/PW1100G)")
        ac_type_spec = st.text_input("On A/C Type/ Speciality:", value="A320/A321 CAT A ( CFM56/PW1100G)")
        
        experience_years = st.number_input("Aircraft/ Engines/ maintenance experience (in years):", min_value=0, value=8)

        st.divider()
        st.markdown("##### 🔍 2. Situation & Request Details")
        background = st.text_area("Situation/ Background:", placeholder="VD: FQI IN DEGRADED MODE")
        request_functions = st.text_area("One-Off Authorization Requested with Functions:", 
                                        placeholder="VD: RAISE MEL 28-07-01-02 FOR A/C A699 IN VCL ON 07FEB26")

        st.divider()
        st.markdown("##### ✈️ 3. Aircraft & Engine Details")
        c1, c2, c3 = st.columns(3)
        ac_type_input = c1.text_input("On A/C type:", value="A320")
        ac_reg = c2.text_input("A/C Reg:", placeholder="VD: VN-A699")
        engine_type = c3.text_input("Engine type:", value="CFM56")
        
        c1, c2 = st.columns(2)
        eng1_sn = c1.text_input("Eng # 1 S/N:", placeholder="VD: 569959")
        eng2_sn = c2.text_input("Eng # 2 S/N:", placeholder="VD: 699161")
        
        c1, c2 = st.columns(2)
        station = c1.text_input("Station:", value="DAD")
        date_duration = c2.text_input("Date/Duration:", value=now_vn.strftime("%d/%m/%Y"))

        st.divider()
        st.markdown("##### 🖊️ 4. Manager Confirmation")
        manager_name = st.text_input("Manager Name:", value="HỒ HỮU ĐÔNG")
        
        submitted = st.form_submit_button("📄 Điền ONE-OFF Template & Tải xuống", type="primary", use_container_width=True)

    if submitted:
        template_file = "one_off.docx"
        if not os.path.exists(template_file):
            st.error(f"❌ Không tìm thấy file mẫu: `{template_file}`. Vui lòng tải file này lên thư mục dự án.")
        else:
            try:
                doc = DocxTemplate(template_file)
                context = {
                    "applicant_name": applicant_name.upper(),
                    "auth_scope": auth_scope,
                    "auth_no": auth_no.upper(),
                    "licence_no": licence_no,
                    "ac_type_spec": ac_type_spec,
                    "expiry_date": expiry_date.upper(),
                    "experience_years": experience_years,
                    "background": background.upper(),
                    "request_functions": request_functions.upper(),
                    "ac_type": ac_type_input.upper(),
                    "ac_reg": ac_reg.upper(),
                    "engine_type": engine_type.upper(),
                    "eng1_sn": eng1_sn,
                    "eng2_sn": eng2_sn,
                    "station": station.upper(),
                    "date_duration": date_duration,
                    "manager_name": manager_name.upper(),
                    "manager_date": date_duration
                }
                doc.render(context)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.success("✅ Đã tạo ONE-OFF Authorization thành công!")
                st.download_button(
                    label="📥 Tải ONE-OFF (.docx)",
                    data=buffer,
                    file_name=f"ONE_OFF_{ac_reg.upper()}_{datetime.now().strftime('%d%m')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"⚠️ Lỗi khi xử lý template: {e}")
